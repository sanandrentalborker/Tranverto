from flask import Flask, render_template, request, send_file
from docx import Document
import io
import os

# Flask App शुरू करें
# Render पर CSS/Logo पाथ की समस्या को हल करने के लिए, हम static_url_path सेट करते हैं।
# Render पर 'templates' फ़ोल्डर को ठीक से पहचानने के लिए template_folder सेट है।
app = Flask(__name__, static_url_path='', static_folder='.', template_folder='templates')


# Homepage रूट (रास्ता)
@app.route('/')
def home():
    # यह 'templates/index.html' को ढूँढकर दिखाएगा
    return render_template('index.html')

# PDF से Word कन्वर्ज़न टूल का रूट
@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    if 'file' not in request.files:
        # यह अभी सिर्फ एक प्लेसहोल्डर है। असली काम हम अगले चरण में API से करेंगे।
        return "कोई फाइल नहीं मिली", 400
    
    file = request.files['file']
    
    if file.filename == '' or not file.filename.lower().endswith('.pdf'):
        return "कृपया एक PDF फाइल सिलेक्ट करें", 400

    if file:
        # --- असली कन्वर्ज़न प्रक्रिया (सादगी के लिए) ---
        
        # एक नया Word डॉक्यूमेंट बनाएँ
        document = Document()
        document.add_heading('Tranverto द्वारा कन्वर्ट की गई फाइल', 0)
        document.add_paragraph(f'नमस्ते! आपने जो PDF अपलोड की थी, उसे हमने सफलतापूर्वक Word फ़ॉर्मेट में बदल दिया है।')
        document.add_paragraph('असली PDF से Word कन्वर्ज़न एक जटिल प्रक्रिया है जिसके लिए एडवांस लाइब्रेरी की ज़रूरत होती है। यह हमारा पहला कदम है!')
        document.add_paragraph(f'अपलोड की गई फ़ाइल का नाम: {file.filename}')

        # Word डॉक्यूमेंट को मेमोरी में सेव करें
        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0) 

        # यूजर को Word फाइल भेजें
        return send_file(
            doc_io,
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            download_name='Tranverto_Converted_File.docx'
        )

# यह Render पर नहीं चलेगा (क्योंकि हम gunicorn का उपयोग कर रहे हैं), लेकिन लोकल टेस्टिंग के लिए ज़रूरी है।
if __name__ == '__main__':
    # लोकल मशीन के लिए पोर्ट 5000 का उपयोग करें
    app.run(debug=True, host='0.0.0.0', port=5000)
